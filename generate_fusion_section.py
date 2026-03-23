#!/usr/bin/env python3
"""Generate the LiDAR-Camera Fusion section as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Styles ────────────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13) if level == 1 else Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def body_run(doc, parts):
    """parts = list of (text, bold) tuples"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def equation(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(11)
    return p

def spacer(doc):
    p = doc.add_paragraph('')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


# ══════════════════════════════════════════════════════════════════════════════
#  TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('LiDAR–Camera Sensor Fusion', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.name = 'Times New Roman'
title.runs[0].font.size = Pt(16)
title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

subtitle = doc.add_paragraph('Research Paper Section — AVMI UGV Dataset')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  INTRO PARAGRAPH
# ══════════════════════════════════════════════════════════════════════════════
body_run(doc, [
    ('To provide complementary geometric and semantic understanding of the environment, we implement a '
     'LiDAR–camera fusion pipeline that projects 3-D point cloud data onto the 2-D camera image plane. '
     'The fused representation assigns each projected LiDAR point its semantic label — encoded directly '
     'in the point cloud — enabling validation of the terrain segmentation model and geometric grounding '
     'of pixel-wise predictions.', False)
])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SUBSECTION 1: Sensor Configuration
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Sensor Configuration', level=2)

body_run(doc, [
    ('The simulation platform (Unreal Engine) equips the UGV with a semantic LiDAR and an RGB camera '
     'mounted at fixed offsets on the vehicle body. In the UE vehicle frame (X-forward, Y-right, Z-up), '
     'the LiDAR is located at ', False),
    ('(−31.08, 0.00, 211.00) cm', True),
    (' and the camera at ', False),
    ('(55.00, −15.00, 200.00) cm', True),
    (', with a camera pitch of −2°. The sensor separation of approximately 86 cm along the vehicle '
     'longitudinal axis defines the system baseline. The camera intrinsic matrix is:', False),
])

equation(doc, '        | 278.59    0     323.80 |')
equation(doc, '   K =  |   0     334.29  229.20 |')
equation(doc, '        |   0       0       1    |')

body_run(doc, [
    ('where ', False), ('fx = 278.59', True), (' and ', False), ('fy = 334.29', True),
    (' are the focal lengths in pixels, and ', False),
    ('(cx, cy) = (323.80, 229.20)', True),
    (' is the principal point, all derived from the Unreal Engine camera calibration.', False),
])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SUBSECTION 2: Coordinate Transformation
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Coordinate Transformation', level=2)

body_run(doc, [
    ('To project LiDAR points into the camera frame, we derive the extrinsic rotation ', False),
    ('R(L→C)', True),
    (' and translation ', False),
    ('t(L→C)', True),
    (' from the known sensor placements. UE world coordinates (cm, Y-right) are first converted to the '
     'ROS vehicle frame (metres, Y-left):', False),
])

equation(doc, '   p_ROS = [ x/100,  -y/100,  z/100 ]ᵀ')

body_run(doc, [
    ('The fixed rotation from the ROS vehicle frame (X-forward, Y-left, Z-up) to the camera optical '
     'frame (X-right, Y-down, Z-forward) is the matrix:', False),
])

equation(doc, '              |  0  -1   0 |')
equation(doc, '   R_base  =  |  0   0  -1 |')
equation(doc, '              |  1   0   0 |')

body_run(doc, [
    ("The camera's in-vehicle rotation ", False),
    ('R_cam', True),
    (' (pitch −2°) is computed from ZYX Euler angles. '
     'The full LiDAR-to-camera extrinsics are then:', False),
])

equation(doc, '   R(L→C) = R_base · R_cam')
equation(doc, '   t(L→C) = R(L→C) · (p_C − p_L)')

body_run(doc, [
    ('where ', False),
    ('p_C', True),
    (' and ', False),
    ('p_L', True),
    (' are the camera and LiDAR positions converted to the ROS vehicle frame.', False),
])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SUBSECTION 3: Projection and Filtering
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Projection and Filtering', level=2)

body_run(doc, [
    ('Each LiDAR point ', False), ('P_L = [x, y, z]ᵀ', True),
    (' is transformed into the camera frame as:', False),
])

equation(doc, '   P_C = R(L→C) · P_L  +  t(L→C)')

body_run(doc, [
    ('Points with depth ', False),
    ('Z_C < 0.5 m', True),
    (' (behind or too close to the camera) and ', False),
    ('Z_C > 50 m', True),
    (' are discarded. The remaining points are projected onto the image plane via the standard '
     'pinhole camera model:', False),
])

equation(doc, '   u = fx · (X_C / Z_C) + cx')
equation(doc, '   v = fy · (Y_C / Z_C) + cy')

body_run(doc, [
    ('Only points whose pixel coordinates ', False),
    ('(u, v)', True),
    (' fall within the image bounds are retained for rendering and publication.', False),
])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SUBSECTION 4: Semantic Point Cloud Coloring
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Semantic Point Cloud Coloring', level=2)

body_run(doc, [
    ('The Unreal Engine semantic LiDAR encodes a terrain class label for each point as a packed RGB '
     'value stored in the fourth channel of the point cloud (IEEE 754 float32, interpreted as '
     '0xRRGGBB). Three terrain classes are represented, consistent with the six-class taxonomy used '
     'for segmentation training:', False),
])

# Bullet list
for item in [
    ('Rock', 'Red   (R=255, G=0, B=0)'),
    ('Ground / Traversable terrain', 'Brown  (R=139, G=90, B=43)'),
    ('Vegetation / Trees', 'Green  (R=0, G=128, B=0)'),
]:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(item[0] + ': ')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(item[1])
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

body_run(doc, [
    ('The RGB channels are extracted from the packed 32-bit integer via bitwise operations:', False),
])

equation(doc, '   r = (rgb >> 16) & 0xFF')
equation(doc, '   g = (rgb >>  8) & 0xFF')
equation(doc, '   b = (rgb      ) & 0xFF')

body_run(doc, [
    ('Each surviving projected point is rendered on the camera image at ', False),
    ('(u, v)', True),
    (' using its semantic colour, producing a fused overlay that directly visualises the geometric '
     'extent and class identity of each terrain region (Fig. X).', False),
])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SUBSECTION 5: Limitations — Parallax
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Limitations: Parallax-Induced Boundary Displacement', level=2)

body_run(doc, [
    ('Due to the physical separation between the LiDAR and camera origins, a parallax-induced boundary '
     'displacement arises at object edges. The pixel displacement is approximated by:', False),
])

equation(doc, '   Δu ≈ (fx · b) / Z')

body_run(doc, [
    ('where ', False),
    ('b', True),
    (' is the lateral sensor baseline and ', False),
    ('Z', True),
    (' is the object depth. For the AVMI sensor configuration with a lateral offset of 15 cm, '
     'the expected displacement at ', False),
    ('Z = 5 m', True),
    (' is approximately 8 pixels, decreasing inversely with distance. This effect is inherent to '
     'multi-sensor systems with non-coincident mounting positions and is a well-documented '
     'characteristic of LiDAR–camera fusion [1]. Interior points within object regions are correctly '
     'labelled; only boundary pixels are affected.', False),
])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, 'References', level=2)

refs = [
    '[1] A. Geiger, P. Lenz, and R. Urtasun, "Are we ready for autonomous driving? The KITTI vision '
     'benchmark suite," in Proc. IEEE CVPR, 2012.',
]
for r in refs:
    p = doc.add_paragraph(r)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = '/home/pinaka/GANav-offroad/fusion_section.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
